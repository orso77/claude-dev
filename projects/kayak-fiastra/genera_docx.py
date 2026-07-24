"""
Genera il Project Plan DOCX per la Scuola Kayak Lago di Fiastra.
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "project-plan-scuola-kayak-fiastra.docx")

doc = Document()

# ── Stili base ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # blu scuro
    hs.font.name = 'Calibri'

# ── Helper ──────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.add_run(text)

# ═══════════════════════════════════════════════════════════
#  COPERTINA
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROJECT PLAN")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Scuola di Kayak e Centro Escursioni Acquatiche")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lago di Fiastra — Provincia di Macerata (Marche)")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documento predisposto per la partecipazione a bandi pubblici\nPNRR — Regione Marche — Fondi strutturali europei")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Versione 2.6.4.15")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ASD Scuola Kayak Fiastra")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  INDICE
# ═══════════════════════════════════════════════════════════
doc.add_heading('Indice', level=1)
indice_voci = [
    "1. Executive Summary",
    "2. Contesto territoriale e analisi di mercato",
    "3. Descrizione del progetto",
    "4. Soggetti coinvolti e governance",
    "5. Obiettivi generali e specifici",
    "6. Attività previste",
    "7. Dotazione strumentale e attrezzature",
    "8. Piano economico-finanziario",
    "9. Cronoprogramma",
    "10. Sostenibilità ambientale, sociale ed economica",
    "11. Indicatori di risultato (KPI)",
    "12. Coerenza con PNRR e strategie regionali",
    "13. Piano di comunicazione e marketing",
    "14. Analisi SWOT",
    "15. Aspetti normativi e autorizzazioni",
    "16. Gestione dei rischi",
    "17. Allegati",
]
for v in indice_voci:
    p = doc.add_paragraph(v)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    "Il presente progetto prevede la costituzione di una Scuola di Kayak e Centro Escursioni "
    "Acquatiche presso il Lago di Fiastra, nel cuore del Parco Nazionale dei Monti Sibillini, "
    "in provincia di Macerata (Regione Marche)."
)
doc.add_paragraph(
    "L'iniziativa nasce dalla collaborazione tra istruttori federali di kayak, "
    "qualificati dalla Federazione Italiana Canoa Kayak (FICK) e dalla Federazione Italiana "
    "Canoa Turistica (FICT), e la Sib.Bellini Turist Service di Leli Emanuela, concessionaria "
    "dell'area lacuale con accesso diretto al lago."
)
doc.add_paragraph(
    "La scuola persegue un duplice obiettivo strategico:"
)
add_bullet(doc, "Formazione tecnica: ", bold_prefix="")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Formazione tecnica: ")
run.bold = True
p.add_run(
    "corsi strutturati di kayak per principianti e intermedi, con utilizzo di kayak sit-in "
    "finalizzati all'apprendimento della tecnica di pagaiata, sicurezza in acqua e autonomia "
    "nella navigazione lacuale."
)

add_bullet(doc, "Escursionismo acquatico guidato: ", bold_prefix="")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Escursionismo acquatico guidato: ")
run.bold = True
p.add_run(
    "escursioni turistiche sul lago con kayak sit-on-top (flotta di circa 20 unità già "
    "disponibili), sempre sotto la guida di istruttori federali qualificati, per offrire "
    "un'esperienza di turismo sportivo accessibile anche ai non esperti."
)

doc.add_paragraph(
    "Il progetto si inserisce pienamente nelle strategie di sviluppo del turismo sostenibile "
    "e sportivo delle aree interne, in coerenza con le linee di intervento del PNRR "
    "(Missione 1 — Componente 3: Turismo e Cultura) e con le politiche regionali delle Marche "
    "per la valorizzazione dei territori colpiti dal sisma 2016. Il Lago di Fiastra, situato "
    "in un'area a forte vocazione naturalistica ma economicamente fragile, rappresenta il "
    "contesto ideale per un intervento che coniughi sviluppo economico locale, destagionalizzazione "
    "dei flussi turistici e promozione di uno stile di vita attivo e sostenibile."
)

doc.add_paragraph(
    "Il kayak e le discipline acquatiche non motorizzate rappresentano l'emblema dello sport "
    "a impatto ambientale zero: nessun motore, nessuna emissione, nessun inquinamento acustico. "
    "L'intero progetto è concepito secondo principi di piena compatibilità ambientale, nel "
    "rispetto del contesto del Parco Nazionale dei Monti Sibillini."
)

doc.add_paragraph(
    "Il progetto pone inoltre particolare attenzione all'inclusività: tra gli obiettivi a "
    "medio termine è prevista l'introduzione di corsi dedicati a persone con disabilità, "
    "per garantire l'accesso allo sport e alla natura al più ampio "
    "numero possibile di persone."
)

doc.add_paragraph(
    "Il finanziamento richiesto è finalizzato all'acquisto delle attrezzature necessarie "
    "(kayak sit-in, equipaggiamento di sicurezza, materiale didattico) e al sostegno delle "
    "spese di avviamento della scuola nel primo anno di attività."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  2. CONTESTO TERRITORIALE E ANALISI DI MERCATO
# ═══════════════════════════════════════════════════════════
doc.add_heading('2. Contesto territoriale e analisi di mercato', level=1)

doc.add_heading('2.1 Il Lago di Fiastra', level=2)
doc.add_paragraph(
    "Il Lago di Fiastra è un bacino artificiale situato a circa 685 m s.l.m. nel comune di "
    "Fiastra, provincia di Macerata. È il secondo lago più grande delle Marche con una superficie "
    "di circa 1 km² e una capacità di circa 20 milioni di m³. Il lago è inserito nel Parco "
    "Nazionale dei Monti Sibillini ed è circondato da un paesaggio montano di grande pregio "
    "naturalistico."
)
doc.add_paragraph(
    "Le caratteristiche del bacino — acque calme, assenza di traffico motorizzato significativo, "
    "temperature miti in estate, paesaggio incontaminato — lo rendono un luogo ideale per la "
    "pratica del kayak a tutti i livelli, dalla formazione base alle escursioni panoramiche."
)

doc.add_heading('2.2 Il contesto socio-economico', level=2)
doc.add_paragraph(
    "L'area di Fiastra rientra nel cosiddetto \"cratere sismico\" del terremoto del Centro Italia "
    "del 2016-2017. La ricostruzione è ancora in corso e il tessuto economico locale, fortemente "
    "dipendente dal turismo stagionale e dall'agricoltura, necessita di interventi di diversificazione "
    "e rilancio. Il presente progetto intende contribuire concretamente a:"
)
add_bullet(doc, "Creare occupazione locale qualificata in un'area dove il sisma ha drasticamente ridotto le opportunità lavorative")
add_bullet(doc, "Estendere la stagione turistica oltre i mesi di luglio e agosto, contrastando la stagionalità che penalizza le aree interne")
add_bullet(doc, "Attrarre un segmento turistico sportivo e attivo, con maggiore capacità di spesa e permanenza media più lunga")
add_bullet(doc, "Offrire ai giovani del territorio un'opportunità di formazione sportiva qualificata, oggi assente nell'area")
add_bullet(doc, "Valorizzare le risorse naturali del territorio in modo sostenibile, trasformando il lago in un polo di attrattività turistica permanente")
add_bullet(doc, "Rafforzare l'offerta turistica integrata (sport + enogastronomia + natura) contribuendo alla rinascita dell'identità territoriale post-sisma")

doc.add_heading('2.3 Analisi della domanda', level=2)
doc.add_paragraph(
    "Il turismo sportivo e outdoor è in forte crescita in Italia e in Europa. Secondo i dati "
    "ISTAT e UNWTO, il segmento del turismo attivo cresce a tassi del 10-15% annuo, con una "
    "domanda particolarmente forte per:"
)
add_bullet(doc, "Esperienze in natura accessibili anche ai principianti")
add_bullet(doc, "Attività acquatiche in acque interne (laghi, fiumi)")
add_bullet(doc, "Proposte che combinano sport, benessere e scoperta del territorio")
add_bullet(doc, "Offerte family-friendly e adatte a gruppi eterogenei")

doc.add_paragraph(
    "Il Lago di Fiastra attira già un flusso significativo di visitatori nel periodo estivo "
    "(stimati 30.000-50.000 presenze/anno nell'area), ma l'offerta di attività strutturate "
    "sul lago è attualmente limitata. La creazione di una scuola di kayak certificata colma "
    "un vuoto nell'offerta turistica locale."
)

doc.add_heading('2.4 Analisi della concorrenza', level=2)
doc.add_paragraph(
    "Nell'area del Lago di Fiastra non sono attualmente presenti scuole di kayak strutturate "
    "e riconosciute dalla FICK/FICT. Le attività di noleggio kayak esistenti sono limitate al "
    "semplice noleggio senza istruzione, senza guida qualificata e senza programma didattico. "
    "Questo posiziona il progetto in una nicchia attualmente non servita, con un chiaro "
    "vantaggio competitivo derivante dalla qualificazione federale degli istruttori."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  3. DESCRIZIONE DEL PROGETTO
# ═══════════════════════════════════════════════════════════
doc.add_heading('3. Descrizione del progetto', level=1)

doc.add_heading('3.1 Visione e missione', level=2)
doc.add_paragraph(
    "Visione: Diventare il punto di riferimento per il kayak e le attività acquatiche "
    "a impatto zero nell'entroterra marchigiano, contribuendo alla rinascita economica "
    "e turistica dell'area del cratere sismico attraverso lo sport sostenibile e il "
    "contatto con la natura. Dimostrare che lo sviluppo turistico e la tutela ambientale "
    "possono rafforzarsi reciprocamente, creando un modello virtuoso per le aree interne."
)
doc.add_paragraph(
    "Missione: Offrire formazione tecnica di kayak di qualità certificata e esperienze "
    "escursionistiche acquatiche sicure, accessibili e memorabili sul Lago di Fiastra, "
    "in un contesto di accoglienza, inclusività e valorizzazione del territorio. "
    "Promuovere lo sport come strumento di coesione sociale, di attrattività turistica "
    "e di riscatto per un territorio duramente colpito dal sisma."
)

doc.add_heading('3.2 Le due linee di attività', level=2)

doc.add_heading('3.2.1 Linea 1 — Scuola di Kayak (formazione tecnica)', level=3)
doc.add_paragraph(
    "La scuola offre corsi strutturati su kayak sit-in, tenuti da "
    "istruttori federali FICK/FICT. I corsi sono articolati su più livelli:"
)

add_table(doc,
    ["Livello", "Durata", "Contenuti principali"],
    [
        ["Base (Discover)", "2-3 ore\n(lezione singola)", "Introduzione al kayak, imbarco/sbarco, "
         "pagaiata base, sicurezza, recupero"],
        ["Principiante", "3 lezioni\n(6-8 ore totali)", "Pagaiata avanti/indietro, virata, "
         "appoggio, traghetto base, navigazione guidata"],
        ["Intermedio", "5 lezioni\n(12-15 ore totali)", "Pagaiata avanzata, roll base, "
         "navigazione autonoma, meteorologia, cartografia"],
        ["Avanzato", "Su richiesta", "Roll avanzato, navigazione in condizioni variabili, "
         "primo soccorso acquatico"],
    ],
    col_widths=[3, 3, 9]
)

doc.add_paragraph()
doc.add_paragraph(
    "I corsi prevedono un rapporto di 1 istruttore ogni 6 allievi, per garantire "
    "un'attenzione personalizzata e la massima sicurezza durante le lezioni in acqua. "
    "Requisito obbligatorio per la partecipazione è la capacità di nuotare."
)

doc.add_paragraph()
doc.add_paragraph(
    "I corsi seguono il programma didattico FICK/FICT e conducono al rilascio della "
    "certificazione \"Pagaia Azzurra\", il sistema di brevetti progressivi della FICK "
    "(Federazione Italiana Canoa Kayak) articolato su più livelli di competenza (SK1, SK2, SK3, ecc.). "
    "Ogni livello attesta il raggiungimento di abilità tecniche specifiche — dalla pagaiata "
    "base e sicurezza in acque calme (SK1) fino alla navigazione avanzata in condizioni "
    "variabili (livelli superiori). La certificazione Pagaia Azzurra è riconosciuta a livello "
    "federale e costituisce un requisito per l'accesso ai livelli formativi successivi."
)

doc.add_heading('3.2.2 Linea 2 — Escursioni guidate (turismo sportivo)', level=3)
doc.add_paragraph(
    "Le escursioni guidate utilizzano la flotta di circa 20 kayak sit-on-top già disponibili. "
    "Sono pensate per turisti senza esperienza pregressa e offrono un'esperienza di scoperta "
    "del lago in totale sicurezza, sempre sotto la guida di almeno un istruttore federale."
)

add_table(doc,
    ["Tipologia", "Durata", "Descrizione"],
    [
        ["Escursione breve", "1-1,5 ore", "Giro panoramico della parte centrale del lago, "
         "adatto a tutti, famiglie incluse"],
        ["Escursione completa", "2,5-3 ore", "Circumnavigazione del lago con soste per "
         "bagno e spiegazioni naturalistiche"],
        ["Sunset Tour", "1,5-2 ore", "Escursione al tramonto con sosta conviviale "
         "al rientro"],
        ["Escursione didattica\n(scuole/gruppi)", "2-3 ore", "Percorso con contenuti di "
         "educazione ambientale, flora e fauna lacustre"],
    ],
    col_widths=[3, 2.5, 9]
)

doc.add_paragraph()
doc.add_paragraph(
    "Tutte le escursioni prevedono un rapporto minimo di 1 guida ogni 6 kayak, "
    "a garanzia della sicurezza e della qualità dell'esperienza. "
    "Requisito obbligatorio per la partecipazione è la capacità di nuotare."
)

doc.add_paragraph()
doc.add_paragraph(
    "Le escursioni includono: briefing iniziale di sicurezza, fornitura di giubbotto "
    "salvagente omologato, pagaia, kayak sit-on-top, sacca stagna e busta porta cellulare "
    "stagna, accompagnamento da parte di guide qualificate. L'attività è interamente a "
    "propulsione umana, senza alcun impatto acustico o ambientale sul lago e sull'ecosistema."
)

doc.add_heading('3.3 Sede operativa e infrastruttura', level=2)
doc.add_paragraph(
    "La sede operativa della scuola è situata presso la concessione dell'area lacuale "
    "della Sib.Bellini Turist Service di Leli Emanuela, partner del progetto, che dispone di:"
)
add_bullet(doc, "Accesso diretto al lago")
add_bullet(doc, "Spazio per deposito attrezzature (kayak, pagaie, giubbotti)")

doc.add_paragraph(
    "Questa configurazione consente di avviare l'attività con investimenti infrastrutturali "
    "minimi, concentrando le risorse sull'acquisto delle attrezzature e sulla promozione."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  4. SOGGETTI COINVOLTI E GOVERNANCE
# ═══════════════════════════════════════════════════════════
doc.add_heading('4. Soggetti coinvolti e governance', level=1)

doc.add_heading('4.1 Composizione del team', level=2)

add_table(doc,
    ["Ruolo", "N.", "Qualifica", "Responsabilità"],
    [
        ["Istruttori federali FICK/FICT", "", "Brevetto istruttore federale\nFICK e/o FICT",
         "Conduzione corsi, guida escursioni,\nsicurezza in acqua, didattica"],
        ["Sib.Bellini Turist Service\ndi Leli Emanuela", "1", "Concessionaria area lago",
         "Logistica, accoglienza,\ngestione prenotazioni, marketing locale"],
    ],
    col_widths=[3.5, 1, 4.5, 5]
)

doc.add_heading('4.2 Forma giuridica', level=2)
doc.add_paragraph(
    "La forma giuridica scelta per la gestione dell'attività è l'ASD — Associazione Sportiva "
    "Dilettantistica, denominata \"ASD Scuola Kayak Fiastra\", affiliata FICK/FICT."
)
doc.add_paragraph(
    "L'ASD è la forma più idonea per l'attività sportiva e didattica, in quanto consente:"
)
add_bullet(doc, "Affiliazione alle federazioni sportive (FICK/FICT)")
add_bullet(doc, "Regime fiscale agevolato previsto per le ASD")
add_bullet(doc, "Accesso ai bandi pubblici destinati alle associazioni sportive")
add_bullet(doc, "Rilascio di attestati e brevetti federali ai partecipanti dei corsi")
doc.add_paragraph(
    "La collaborazione con la Sib.Bellini Turist Service di Leli Emanuela è regolamentata da un accordo di partenariato "
    "che disciplina l'utilizzo della sede, della logistica e delle attrezzature già disponibili."
)

doc.add_heading('4.3 Governance e organizzazione', level=2)
doc.add_paragraph("Si prevede la seguente struttura organizzativa:")
add_bullet(doc, "Presidente / Legale rappresentante")
add_bullet(doc, "Direttore tecnico — responsabile programma didattico")
add_bullet(doc, "Responsabile logistica e accoglienza")
add_bullet(doc, "Responsabile sicurezza — responsabile protocolli di sicurezza")
add_bullet(doc, "Responsabile marketing e comunicazione")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  5. OBIETTIVI
# ═══════════════════════════════════════════════════════════
doc.add_heading('5. Obiettivi generali e specifici', level=1)

doc.add_heading('5.1 Obiettivi generali', level=2)
add_numbered(doc, "Promuovere la pratica del kayak e delle discipline acquatiche non motorizzate come sport a impatto ambientale zero, accessibile e sostenibile")
add_numbered(doc, "Contribuire alla rinascita economica e turistica dell'area del cratere sismico, diversificando e destagionalizzando l'offerta del Lago di Fiastra")
add_numbered(doc, "Generare occupazione qualificata e offrire opportunità formative ai giovani in un territorio a forte rischio di spopolamento")
add_numbered(doc, "Valorizzare le risorse naturali e paesaggistiche del territorio in modo responsabile, nel pieno rispetto dell'ecosistema del Parco Nazionale")
add_numbered(doc, "Garantire l'inclusività e l'accessibilità dell'attività sportiva, con programmi dedicati a persone con disabilità")
add_numbered(doc, "Creare un modello replicabile di turismo sportivo sostenibile nelle aree interne, capace di coniugare sviluppo economico e tutela ambientale")

doc.add_heading('5.2 Obiettivi specifici (Anno 1)', level=2)
add_table(doc,
    ["Obiettivo", "Target", "Indicatore"],
    [
        ["Corsi di kayak erogati", "≥ 40 corsi (base + principiante)", "N. corsi completati"],
        ["Allievi formati", "≥ 120 persone", "N. attestati rilasciati"],
        ["Escursioni guidate", "≥ 150 escursioni", "N. escursioni effettuate"],
        ["Turisti coinvolti nelle escursioni", "≥ 1.500 persone", "N. partecipanti registrati"],
        ["Periodo di attività", "Maggio — Ottobre (6 mesi)", "Mesi di apertura effettiva"],
        ["Collaborazioni con scuole", "≥ 3 istituti", "N. convenzioni attivate"],
        ["Presenza online", "Sito web + social attivi", "N. follower, recensioni"],
    ]
)

doc.add_heading('5.3 Obiettivi a medio termine (Anno 2-3)', level=2)
add_bullet(doc, "Ampliare la flotta sit-in a 10 unità per gestire corsi paralleli")
add_bullet(doc, "Ottenere il riconoscimento FICK/FICT come centro federale")
add_bullet(doc, "Introdurre corsi per bambini (6-12 anni) con kayak dedicati")
add_bullet(doc, "Introdurre corsi per persone con disabilità")
add_bullet(doc, "Attivare convenzioni con tour operator e piattaforme di booking esperienziale")
add_bullet(doc, "Estendere la stagione ad aprile-novembre con attività per scuole e gruppi organizzati")
add_bullet(doc, "Raggiungere il pareggio di bilancio e iniziare a generare margine positivo")
add_bullet(doc, "Introdurre attività complementari: SUP, canoa canadese, kayak polo")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  6. ATTIVITÀ PREVISTE
# ═══════════════════════════════════════════════════════════
doc.add_heading('6. Attività previste', level=1)

doc.add_heading('6.1 Fase di avviamento (mesi 1-3)', level=2)
add_numbered(doc, "Costituzione dell'ASD Scuola Kayak Fiastra")
add_numbered(doc, "Affiliazione alla Federazione Italiana Canoa Kayak (FICK) e alla Federazione Italiana Canoa Turistica (FICT)")
add_numbered(doc, "Acquisizione delle autorizzazioni necessarie (Comune di Fiastra, Parco Nazionale)")
add_numbered(doc, "Stipula dell'accordo di partenariato con la Sib.Bellini Turist Service di Leli Emanuela")
add_numbered(doc, "Acquisto attrezzature (kayak sit-in, pagaie, giubbotti, equipaggiamento sicurezza)")
add_numbered(doc, "Allestimento area deposito e punto di imbarco/sbarco")
add_numbered(doc, "Redazione del programma didattico e dei protocolli di sicurezza")
add_numbered(doc, "Creazione del sito web e dei canali social")
add_numbered(doc, "Stipula delle assicurazioni obbligatorie (RC, infortuni)")
add_numbered(doc, "Sottoscrizione convenzioni con strutture ricettive locali")

doc.add_heading('6.2 Fase operativa (mesi 4-9 — stagione maggio-ottobre)', level=2)
add_numbered(doc, "Apertura della scuola e avvio dei corsi di kayak")
add_numbered(doc, "Avvio delle escursioni guidate sul lago")
add_numbered(doc, "Attività di promozione e marketing continua (social, sito, eventi)")
add_numbered(doc, "Partecipazione a eventi locali e fiere del turismo")
add_numbered(doc, "Monitoraggio della soddisfazione dei partecipanti (questionari)")
add_numbered(doc, "Raccolta dati per rendicontazione bando")

doc.add_heading('6.3 Fase di consolidamento (mesi 10-12)', level=2)
add_numbered(doc, "Chiusura stagionale e messa in sicurezza delle attrezzature")
add_numbered(doc, "Analisi dei risultati della prima stagione")
add_numbered(doc, "Rendicontazione del bando (se finanziato)")
add_numbered(doc, "Pianificazione della stagione successiva")
add_numbered(doc, "Formazione continua degli istruttori (aggiornamenti FICK/FICT)")
add_numbered(doc, "Manutenzione e inventario attrezzature")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  7. DOTAZIONE STRUMENTALE E ATTREZZATURE
# ═══════════════════════════════════════════════════════════
doc.add_heading('7. Dotazione strumentale e attrezzature', level=1)

doc.add_heading('7.1 Attrezzature già disponibili', level=2)
add_table(doc,
    ["Attrezzatura", "Quantità", "Stato", "Note"],
    [
        ["Kayak sit-on-top\n(singoli e doppi)", "~20", "Disponibili", "Flotta della Sib.Bellini Turist Service di Leli Emanuela,\ncapacità fino a ~30 persone"],
        ["Pagaie per sit-on-top", "~20", "Disponibili", "In dotazione con i kayak"],
        ["Area deposito", "1", "Disponibile", "Presso la sede della Sib.Bellini Turist Service di Leli Emanuela"],
        ["Punto imbarco/sbarco", "1", "Disponibile", "In concessione dell'area lacuale"],
    ]
)

doc.add_heading('7.2 Attrezzature da acquistare', level=2)
add_table(doc,
    ["Attrezzatura", "Quantità", "Prezzo unitario (stima)", "Totale (stima)", "Note"],
    [
        ["Kayak sit-in 450 cm\n(tipo Oasis 450 o equivalente)", "6", "€ 800 - 1.200", "€ 4.800 - 7.200",
         "Per corsi di formazione tecnica"],
        ["Pagaie per sit-in\n(misure assortite)", "8", "€ 60 - 120", "€ 480 - 960",
         "6 in uso + 2 di scorta"],
        ["Pagaie per sit-on-top\n(integrazione)", "14", "€ 40 - 80", "€ 560 - 1.120",
         "Per raggiungere 30 posti\n(~20 già disponibili + 14 nuove)"],
        ["Paraspruzzi (spraydeck)", "8", "€ 40 - 80", "€ 320 - 640",
         "Per kayak sit-in"],
        ["Giubbotti salvagente\nomologati (CE EN 12402)", "40", "€ 40 - 80", "€ 1.600 - 3.200",
         "Taglie assortite, per tutti\ni partecipanti (36 posti + scorta)"],
        ["Caschi da kayak", "8", "€ 30 - 50", "€ 240 - 400",
         "Per i corsi avanzati"],
        ["Sacche stagne", "40", "€ 15 - 30", "€ 600 - 1.200",
         "Per effetti personali, una per\npartecipante (36 posti + scorta)"],
        ["Buste porta cellulare\nstagne", "40", "€ 8 - 15", "€ 320 - 600",
         "Una per partecipante\n(36 posti + scorta)"],
        ["Kit primo soccorso\nwaterproof", "4", "€ 40 - 60", "€ 160 - 240",
         "Uno per ogni guida in attività"],
        ["Corde di sicurezza\ne throw bag", "4", "€ 25 - 40", "€ 100 - 160",
         "Una per ogni guida in attività"],
        ["Radio VHF impermeabili", "4", "€ 50 - 80", "€ 200 - 320",
         "Una per ogni guida in attività"],
        ["Carrelli porta-kayak", "2", "€ 80 - 150", "€ 160 - 300",
         "Per movimentazione a terra"],
        ["Segnaletica e materiale\ndidattico", "1 kit", "€ 300 - 500", "€ 300 - 500",
         "Cartelli, manuali, poster"],
        ["Abbigliamento tecnico\nistruttori (mute, gilet)", "4 kit", "€ 200 - 350", "€ 800 - 1.400",
         "Muta shorty + gilet per ogni guida"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Totale stimato attrezzature da acquistare: € 10.640 - 18.240")
run.bold = True

doc.add_paragraph(
    "Le stime sono indicative e basate su prezzi di mercato al 2026. I costi effettivi "
    "saranno definiti in fase di approvvigionamento con richiesta di preventivi a fornitori "
    "specializzati."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  8. PIANO ECONOMICO-FINANZIARIO
# ═══════════════════════════════════════════════════════════
doc.add_heading('8. Piano economico-finanziario', level=1)

doc.add_heading('8.1 Investimento iniziale', level=2)
add_table(doc,
    ["Voce di costo", "Importo minimo", "Importo massimo", "Note"],
    [
        ["Attrezzature sportive\n(vedi sezione 7)", "€ 10.640", "€ 18.240", "Kayak, pagaie, sicurezza, ecc."],
        ["Costituzione ASD", "€ 500", "€ 2.000", "Notaio, registrazione, affiliazione"],
        ["Affiliazione FICK/FICT\ne tesseramento", "€ 300", "€ 500", "Quota annuale + tessere istruttori"],
        ["Assicurazioni\n(RC + infortuni)", "€ 1.500", "€ 3.000", "Polizza annuale per attività sportiva"],
        ["Sito web e\nmateriale marketing", "€ 800", "€ 2.000", "Sito, grafica, stampa depliant"],
        ["Allestimento area\ndeposito/imbarco", "€ 500", "€ 1.500", "Rastrelliere, segnaletica, piccoli lavori"],
        ["Spese burocratiche\ne autorizzazioni", "€ 300", "€ 800", "Pratiche Comune, Parco"],
        ["Fondo imprevisti (10%)", "€ 1.454", "€ 2.804", ""],
        ["TOTALE INVESTIMENTO", "€ 15.994", "€ 30.844", ""],
    ]
)

doc.add_heading('8.2 Costi di gestione annuali (stima Anno 1)', level=2)
add_table(doc,
    ["Voce", "Importo annuo (stima)", "Note"],
    [
        ["Assicurazioni (rinnovo)", "€ 1.500 - 3.000", "RC + infortuni"],
        ["Affiliazione FICK/FICT (rinnovo)", "€ 300 - 500", ""],
        ["Manutenzione attrezzature", "€ 500 - 1.000", "Riparazioni, sostituzioni minori"],
        ["Marketing e promozione", "€ 500 - 1.500", "Social ads, eventi, materiale"],
        ["Utenze e consumabili", "€ 300 - 500", ""],
        ["Commercialista / contabilità", "€ 500 - 1.000", "Se necessario"],
    ]
)

doc.add_heading('8.3 Ricavi previsti (stima Anno 1)', level=2)
add_table(doc,
    ["Fonte di ricavo", "Volume stimato", "Prezzo medio", "Ricavo stimato"],
    [
        ["Corsi base (singola lezione)", "80 partecipanti", "€ 40 - 50", "€ 3.200 - 4.000"],
        ["Corsi principiante (3 lezioni)", "30 partecipanti", "€ 100 - 130", "€ 3.000 - 3.900"],
        ["Corsi intermedio (5 lezioni)", "10 partecipanti", "€ 180 - 220", "€ 1.800 - 2.200"],
        ["Escursione breve (1h)", "600 partecipanti", "€ 20 - 25", "€ 12.000 - 15.000"],
        ["Escursione completa (3h)", "500 partecipanti", "€ 35 - 45", "€ 17.500 - 22.500"],
        ["Sunset Tour", "200 partecipanti", "€ 30 - 40", "€ 6.000 - 8.000"],
        ["Escursioni scolastiche", "200 partecipanti", "€ 15 - 20", "€ 3.000 - 4.000"],
        ["TOTALE RICAVI ANNO 1", "", "", "€ 46.500 - 59.600"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Nota: I ricavi sono stimati in modo prudenziale. L'effettivo andamento dipenderà "
    "dalla stagionalità, dalle condizioni meteo e dall'efficacia del marketing."
)

doc.add_heading('8.4 Punto di pareggio (break-even)', level=2)
doc.add_paragraph(
    "Con costi fissi annuali stimati tra € 3.600 e € 7.500 (escludendo l'investimento "
    "iniziale coperto dal finanziamento) e un margine medio per partecipante di € 20-30, "
    "il punto di pareggio operativo si raggiunge con circa 180-375 partecipanti/anno, "
    "target ampiamente alla portata delle stime conservative. L'investimento iniziale, "
    "se coperto dal finanziamento pubblico, non grava sul conto economico."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  9. CRONOPROGRAMMA
# ═══════════════════════════════════════════════════════════
doc.add_heading('9. Cronoprogramma', level=1)

doc.add_paragraph(
    "Il cronoprogramma prevede un arco temporale di 12 mesi dalla data di avvio del progetto "
    "(indicativamente da gennaio/febbraio a dicembre dello stesso anno)."
)

add_table(doc,
    ["Attività", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "M8", "M9", "M10", "M11", "M12"],
    [
        ["Costituzione soggetto giuridico", "■", "■", "", "", "", "", "", "", "", "", "", ""],
        ["Affiliazione FICK/FICT", "", "■", "■", "", "", "", "", "", "", "", "", ""],
        ["Autorizzazioni e permessi", "■", "■", "■", "", "", "", "", "", "", "", "", ""],
        ["Accordo di partenariato", "■", "■", "", "", "", "", "", "", "", "", "", ""],
        ["Acquisto attrezzature", "", "■", "■", "", "", "", "", "", "", "", "", ""],
        ["Allestimento sede operativa", "", "", "■", "■", "", "", "", "", "", "", "", ""],
        ["Sito web e social media", "", "■", "■", "■", "", "", "", "", "", "", "", ""],
        ["Assicurazioni", "", "", "■", "", "", "", "", "", "", "", "", ""],
        ["Protocolli sicurezza e didattica", "", "■", "■", "", "", "", "", "", "", "", "", ""],
        ["APERTURA SCUOLA", "", "", "", "", "■", "", "", "", "", "", "", ""],
        ["Corsi di kayak", "", "", "", "", "■", "■", "■", "■", "■", "■", "", ""],
        ["Escursioni guidate", "", "", "", "", "■", "■", "■", "■", "■", "■", "", ""],
        ["Marketing e promozione", "", "", "■", "■", "■", "■", "■", "■", "■", "■", "", ""],
        ["Monitoraggio e raccolta dati", "", "", "", "", "■", "■", "■", "■", "■", "■", "", ""],
        ["Chiusura stagionale", "", "", "", "", "", "", "", "", "", "■", "", ""],
        ["Rendicontazione e bilancio", "", "", "", "", "", "", "", "", "", "", "■", "■"],
        ["Pianificazione anno successivo", "", "", "", "", "", "", "", "", "", "", "■", "■"],
    ],
    col_widths=[4.5] + [1.05] * 12
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  10. SOSTENIBILITÀ
# ═══════════════════════════════════════════════════════════
doc.add_heading('10. Sostenibilità ambientale, sociale ed economica', level=1)

doc.add_heading('10.1 Sostenibilità ambientale — sport a impatto zero', level=2)
doc.add_paragraph(
    "Il kayak, il SUP e le discipline acquatiche non motorizzate sono tra le attività sportive "
    "a più basso impatto ambientale in assoluto. L'intero progetto è concepito per essere "
    "pienamente compatibile con l'ecosistema del Parco Nazionale dei Monti Sibillini e con "
    "la tutela delle acque del lago."
)
add_bullet(doc, "Impatto ambientale zero: nessun motore, nessuna emissione di CO₂, nessun inquinamento acustico, nessuno scarico in acqua")
add_bullet(doc, "Nessuna infrastruttura invasiva: l'attività si svolge su strutture già esistenti, senza consumo di suolo né alterazione del paesaggio")
add_bullet(doc, "Le escursioni diventano strumento di educazione ambientale: ogni uscita sul lago è un'occasione per conoscere e rispettare l'ecosistema lacustre, la fauna e la flora del Parco")
add_bullet(doc, "Adozione rigorosa dei principi Leave No Trace in tutte le attività, con formazione specifica degli istruttori")
add_bullet(doc, "Sensibilizzazione attiva dei partecipanti sulla tutela delle acque, della biodiversità e della fragilità degli ambienti montani")
add_bullet(doc, "Utilizzo di materiali e attrezzature durevoli, con manutenzione programmata per ridurre i rifiuti e allungare il ciclo di vita dei prodotti")
add_bullet(doc, "Promozione di una cultura del turismo responsabile: il visitatore non è semplice consumatore ma custode temporaneo del territorio")

doc.add_heading('10.2 Sostenibilità sociale e inclusività', level=2)
add_bullet(doc, "Creazione di occupazione qualificata in un'area economicamente fragile, duramente colpita dal sisma 2016")
add_bullet(doc, "Attività inclusiva e accessibile a tutte le fasce d'età (bambini, adulti, anziani), senza barriere di genere o condizione fisica")
add_bullet(doc, "Programmi dedicati a persone con disabilità motoria e cognitiva, per garantire il diritto allo sport e al contatto con la natura a tutti")
add_bullet(doc, "Opportunità formativa e sportiva per i giovani del territorio, oggi privo di strutture sportive acquatiche qualificate")
add_bullet(doc, "Coinvolgimento delle scuole locali con programmi di educazione ambientale e sportiva, contribuendo alla formazione delle nuove generazioni")
add_bullet(doc, "Rafforzamento del senso di comunità e identità locale attraverso lo sport, elemento fondamentale per la ricostruzione del tessuto sociale post-sisma")
add_bullet(doc, "Contributo concreto alla lotta contro lo spopolamento delle aree interne, offrendo ragioni per restare e motivi per tornare")

doc.add_heading('10.3 Sostenibilità economica e sviluppo territoriale', level=2)
add_bullet(doc, "Modello a costi fissi contenuti grazie alla partnership con la Sib.Bellini Turist Service di Leli Emanuela (sede, logistica)")
add_bullet(doc, "Duplice fonte di ricavo (corsi + escursioni) che diversifica il rischio e stabilizza i flussi economici")
add_bullet(doc, "Stagione operativa estendibile (maggio-ottobre, potenzialmente aprile-novembre), contribuendo alla destagionalizzazione turistica dell'area")
add_bullet(doc, "Potenziale di crescita con servizi aggiuntivi a impatto zero (SUP, canoa, corsi per disabili, eventi aziendali)")
add_bullet(doc, "Break-even raggiungibile già al primo anno con volumi conservativi")
add_bullet(doc, "Effetto moltiplicatore sul territorio: ogni visitatore attratto dalla scuola genera ricadute economiche su ricettività, commercio e servizi locali")
add_bullet(doc, "La scuola come polo di attrattività permanente per il Lago di Fiastra, non legato a eventi occasionali ma a un'offerta strutturata e continuativa")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  11. KPI
# ═══════════════════════════════════════════════════════════
doc.add_heading('11. Indicatori di risultato (KPI)', level=1)

add_table(doc,
    ["KPI", "Target Anno 1", "Target Anno 2", "Target Anno 3", "Metodo di rilevazione"],
    [
        ["N. corsi erogati", "≥ 40", "≥ 60", "≥ 80", "Registro attività"],
        ["N. allievi formati", "≥ 120", "≥ 200", "≥ 300", "Registro presenze + attestati"],
        ["N. escursioni effettuate", "≥ 150", "≥ 200", "≥ 250", "Registro attività"],
        ["N. partecipanti escursioni", "≥ 1.500", "≥ 2.500", "≥ 3.500", "Registro presenze"],
        ["Tasso di soddisfazione", "≥ 85%", "≥ 90%", "≥ 92%", "Questionari post-attività"],
        ["Recensioni online (media)", "≥ 4.2/5", "≥ 4.5/5", "≥ 4.5/5", "Google / TripAdvisor"],
        ["Ricavi lordi", "€ 46.500+", "€ 65.000+", "€ 80.000+", "Contabilità"],
        ["Convenzioni scuole", "≥ 3", "≥ 5", "≥ 8", "Contratti firmati"],
        ["Follower social", "≥ 1.000", "≥ 3.000", "≥ 5.000", "Analytics social"],
        ["Occupazione generata", "Avvio attività", "Consolidamento", "Espansione", "Contratti / incarichi"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  12. COERENZA CON PNRR E STRATEGIE REGIONALI
# ═══════════════════════════════════════════════════════════
doc.add_heading('12. Coerenza con PNRR e strategie regionali', level=1)

doc.add_heading('12.1 Piano Nazionale di Ripresa e Resilienza (PNRR)', level=2)
doc.add_paragraph("Il progetto è coerente con le seguenti componenti del PNRR:")

p = doc.add_paragraph()
run = p.add_run("Missione 1 — Componente 3: Turismo e Cultura 4.0")
run.bold = True

add_bullet(doc, "Investimento 4.1: Hub del turismo digitale — il progetto prevede digitalizzazione "
           "della promozione e prenotazione online")
add_bullet(doc, "Investimento 4.2: Fondi integrati per la competitività delle imprese turistiche — "
           "acquisto attrezzature e ammodernamento dell'offerta")

p = doc.add_paragraph()
run = p.add_run("Missione 5 — Componente 2: Rigenerazione urbana e inclusione sociale")
run.bold = True

add_bullet(doc, "Investimento 2.1: Investimenti in progetti di rigenerazione urbana per le aree "
           "del cratere sismico — rilancio economico attraverso turismo sportivo")
add_bullet(doc, "Interventi per le aree interne — valorizzazione dei borghi e dei territori "
           "montani e lacustri")

p = doc.add_paragraph()
run = p.add_run("Missione 5 — Componente 3: Interventi speciali per la coesione territoriale")
run.bold = True

add_bullet(doc, "Strategia Nazionale Aree Interne (SNAI) — il progetto è localizzato in un'area "
           "interna classificata e contribuisce alla riduzione del divario territoriale")

doc.add_heading('12.2 Strategie regionali — Regione Marche', level=2)
doc.add_paragraph("Il progetto si allinea a:")

add_bullet(doc, "POR FESR Marche 2021-2027: Asse prioritario competitività e innovazione delle imprese")
add_bullet(doc, "PSR Marche 2023-2027: misure per la diversificazione economica delle aree rurali")
add_bullet(doc, "Legge Regionale 9/2006: promozione dell'attività sportiva e impiantistica")
add_bullet(doc, "Piano regionale per il turismo sostenibile nelle aree del cratere sismico")
add_bullet(doc, "Strategia regionale per lo sport all'aria aperta e il turismo attivo")
add_bullet(doc, "Bandi regionali per ASD: acquisto attrezzature sportive, avvio nuove attività")

doc.add_heading('12.3 Principi trasversali PNRR', level=2)
add_table(doc,
    ["Principio", "Coerenza del progetto"],
    [
        ["DNSH\n(Do No Significant Harm)", "Piena conformità: sport a impatto ambientale zero, "
         "nessuna emissione, nessun motore, nessuna infrastruttura invasiva, "
         "nessuna alterazione del paesaggio. Il kayak è tra le discipline "
         "sportive più compatibili con il principio DNSH"],
        ["Parità di genere", "Team misto (Sib.Bellini Turist Service di Leli Emanuela "
         "come partner fondatore), attività aperta a tutti i generi. "
         "Programmi inclusivi senza barriere di genere"],
        ["Giovani", "Target privilegiato: formazione sportiva certificata (Pagaia Azzurra) "
         "per i giovani del cratere sismico, oggi privi di offerta sportiva "
         "acquatica. Tirocini e collaborazioni con università e scuole"],
        ["Mezzogiorno / Aree interne", "Progetto nel cuore del cratere sismico 2016. "
         "Contributo diretto alla rinascita economica, al contrasto dello "
         "spopolamento e alla coesione territoriale"],
        ["Inclusività", "Programmi dedicati a persone con disabilità, "
         "accessibilità a tutte le fasce d'età e condizioni fisiche. "
         "Diritto allo sport e alla natura per tutti"],
        ["Digitalizzazione", "Prenotazione online, marketing digitale, "
         "pagamenti elettronici, CRM per gestione clienti"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  13. PIANO DI COMUNICAZIONE E MARKETING
# ═══════════════════════════════════════════════════════════
doc.add_heading('13. Piano di comunicazione e marketing', level=1)

doc.add_heading('13.1 Identità e posizionamento', level=2)
doc.add_paragraph(
    "La scuola si posiziona come l'unica scuola di kayak certificata FICK/FICT sul Lago di Fiastra, "
    "con un'offerta che combina formazione sportiva di qualità ed esperienze turistiche uniche "
    "in un contesto naturalistico di pregio."
)
doc.add_paragraph("Claim suggerito: \"Pagaia, scopri, vivi il lago\" (da validare)")

doc.add_heading('13.2 Canali di comunicazione', level=2)
add_table(doc,
    ["Canale", "Azione", "Frequenza", "Budget stimato"],
    [
        ["Sito web", "Creazione sito con booking online,\ndescrizione corsi, galleria foto", "Continuo", "€ 500 - 1.500"],
        ["Instagram", "Contenuti foto/video del lago e delle\nattività, stories, reels", "3-5 post/settimana", "€ 200 - 500 (ads)"],
        ["Facebook", "Pagina + gruppo locale,\neventi, recensioni", "3-4 post/settimana", "€ 100 - 300 (ads)"],
        ["TripAdvisor /\nGoogle Business", "Profilo attività, raccolta recensioni", "Continuo", "Gratuito"],
        ["Portali booking\nesperienziale", "Inserimento su GetYourGuide,\nViator, Civitatis, Airbnb Experiences", "Continuo", "Commissione"],
        ["Stampa locale", "Comunicati stampa, articoli,\ninterviste", "Al lancio + eventi", "€ 100 - 300"],
        ["Materiale cartaceo", "Depliant, locandine presso hotel,\nB&B, IAT, bar della zona", "Inizio stagione", "€ 200 - 500"],
        ["Passaparola e\nrecensioni", "Incentivo con sconto per\nchi porta amici", "Continuo", "Incluso"],
    ]
)

doc.add_heading('13.3 Partnership e collaborazioni', level=2)
add_bullet(doc, "Strutture ricettive locali (B&B, agriturismi, hotel): pacchetti combinati pernottamento + kayak")
add_bullet(doc, "Enti del turismo (IAT, Pro Loco, Comune di Fiastra): promozione istituzionale")
add_bullet(doc, "Tour operator incoming Marche: inserimento nei pacchetti di turismo attivo")
add_bullet(doc, "Scuole e centri estivi: tariffe convenzionate per attività didattiche")
add_bullet(doc, "Associazioni escursionistiche (CAI, gruppi trekking): escursioni combinate terra-acqua")
add_bullet(doc, "Influencer e content creator del settore outdoor: inviti per esperienza gratuita in cambio di visibilità")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  14. ANALISI SWOT
# ═══════════════════════════════════════════════════════════
doc.add_heading('14. Analisi SWOT', level=1)

add_table(doc,
    ["", "Fattori positivi", "Fattori negativi"],
    [
        ["Fattori\ninterni",
         "PUNTI DI FORZA\n"
         "• Istruttori federali certificati FICK/FICT\n"
         "• Sport a impatto ambientale zero\n"
         "• Sede con concessione dell'area lacuale\n"
         "• Flotta sit-on-top già disponibile\n"
         "• Nessun concorrente diretto\n"
         "• Vocazione inclusiva (disabilità)\n"
         "• Contesto paesaggistico unico\n"
         "• Sinergia con attività turistica locale",
         "PUNTI DI DEBOLEZZA\n"
         "• Nessuna struttura giuridica esistente\n"
         "• Distanza degli istruttori dal lago\n"
         "• Stagionalità dell'attività\n"
         "• Necessità di acquistare kayak sit-in\n"
         "• Brand non ancora conosciuto\n"
         "• Dipendenza dalle condizioni meteo"],
        ["Fattori\nesterni",
         "OPPORTUNITÀ\n"
         "• Fondi PNRR e regionali per cratere sismico\n"
         "• Turismo outdoor e green in forte crescita\n"
         "• Crescente attenzione a sport inclusivi\n"
         "• Area del cratere = priorità per bandi\n"
         "• Nessuna scuola kayak nella zona\n"
         "• Domanda di esperienze in natura\n"
         "• Espansione con SUP, canoa, corsi disabili\n"
         "• Digital marketing a basso costo",
         "MINACCE\n"
         "• Burocrazia e tempi autorizzazioni\n"
         "• Variazioni livello acqua del lago\n"
         "• Eventuali restrizioni del Parco Nazionale\n"
         "• Ingresso di nuovi concorrenti\n"
         "• Condizioni meteo sfavorevoli\n"
         "• Incertezza sui tempi dei finanziamenti\n"
         "• Calo generale dei flussi turistici"],
    ],
    col_widths=[2.5, 6.5, 6.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  15. ASPETTI NORMATIVI E AUTORIZZAZIONI
# ═══════════════════════════════════════════════════════════
doc.add_heading('15. Aspetti normativi e autorizzazioni', level=1)

doc.add_heading('15.1 Autorizzazioni necessarie', level=2)
add_table(doc,
    ["Autorizzazione / Adempimento", "Ente competente", "Stato", "Note"],
    [
        ["Costituzione ASD", "Agenzia Entrate / Notaio", "Da fare",
         "Registrazione e codice fiscale"],
        ["Affiliazione FICK/FICT", "Federazione Italiana\nCanoa Kayak e Canoa Turistica", "Da fare",
         "Richiede statuto conforme e istruttori tesserati"],
        ["Iscrizione Registro CONI /\nSport e Salute", "CONI / Sport e Salute", "Da fare",
         "Obbligatoria per ASD riconosciuta"],
        ["Autorizzazione uso invaso\n(attività sportiva)", "ENEL (proprietario invaso)", "Disponibile",
         "Il Lago di Fiastra è un invaso artificiale\ndi proprietà ENEL. Autorizzazione già ottenuta"],
        ["Permessi per strutture\n(pontili, piattaforme, ecc.)", "Comune di Fiastra", "Da verificare",
         "Permessi urbanistici per eventuali\nstrutture fisse o galleggianti"],
        ["Nulla osta Parco Nazionale\ndei Monti Sibillini", "Ente Parco Nazionale", "Da richiedere",
         "Per attività organizzate in area Parco"],
        ["SCIA / Comunicazione\navvio attività", "Comune di Fiastra\n(SUAP)", "Da fare",
         "Se richiesto per l'attività specifica"],
        ["Assicurazione RC\nverso terzi", "Compagnia assicurativa", "Da stipulare",
         "Obbligatoria, copertura attività sportiva\nacquatica + RC struttura"],
        ["Assicurazione infortuni\npartecipanti", "Compagnia assicurativa", "Da stipulare",
         "Inclusa nel tesseramento FICK/FICT o integrativa"],
        ["Certificazione attrezzature\n(giubbotti CE)", "Fornitore / autocertificazione", "Da verificare",
         "Giubbotti conformi EN 12402,\nkayak conformi"],
        ["Regolamento navigazione\nlacuale", "ENEL", "Disponibile",
         "Regolamento di navigazione sull'invaso\nstabilito da ENEL. Già acquisito"],
        ["Privacy / GDPR", "Garante Privacy", "Da adempiere",
         "Informativa per raccolta dati partecipanti,\nnomina DPO se necessario"],
    ]
)

doc.add_heading('15.2 Normativa di riferimento', level=2)
add_bullet(doc, "D.Lgs. 36/2021 (Riforma dello sport) — disciplina ASD")
add_bullet(doc, "Regolamento FICK/FICT per scuole e istruttori")
add_bullet(doc, "Codice della navigazione (navigazione acque interne)")
add_bullet(doc, "D.Lgs. 81/2008 (Sicurezza sul lavoro) — per quanto applicabile")
add_bullet(doc, "Regolamento del Parco Nazionale dei Monti Sibillini")
add_bullet(doc, "Normativa regionale Marche in materia di turismo e sport")
add_bullet(doc, "Regolamento UE 2016/679 (GDPR)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  16. GESTIONE DEI RISCHI
# ═══════════════════════════════════════════════════════════
doc.add_heading('16. Gestione dei rischi', level=1)

add_table(doc,
    ["Rischio", "Probabilità", "Impatto", "Mitigazione"],
    [
        ["Ritardo nelle autorizzazioni", "Media", "Alto",
         "Avviare le pratiche con largo anticipo;\nconsulente burocratico dedicato"],
        ["Condizioni meteo avverse\nprolungate", "Media", "Medio",
         "Programma flessibile con recuperi;\ncalendario buffer; attività indoor alternative\n(teoria, briefing, team building)"],
        ["Livello acqua lago insufficiente", "Bassa", "Alto",
         "Monitoraggio periodico;\npiani alternativi; comunicazione trasparente"],
        ["Infortunio partecipante", "Bassa", "Alto",
         "Protocolli sicurezza rigorosi;\nassicurazione adeguata; kit primo soccorso;\nformazione BLS istruttori"],
        ["Scarsa affluenza turistica", "Bassa", "Medio",
         "Marketing proattivo; diversificazione\ncanali; pricing flessibile;\npartnership con strutture ricettive"],
        ["Conflitti tra soci / partner", "Bassa", "Alto",
         "Accordo scritto dettagliato;\nstatuto con clausole di risoluzione;\nmediazione preventiva"],
        ["Nuovi concorrenti", "Media", "Basso",
         "Fidelizzazione; qualità certificata FICK/FICT;\nfirst-mover advantage;\nbrand reputation solida"],
        ["Mancato ottenimento\nfinanziamento", "Media", "Alto",
         "Piano B con autofinanziamento\nridotto; partecipazione a più bandi;\ncrowdfunding"],
        ["Danneggiamento attrezzature", "Media", "Basso",
         "Manutenzione programmata;\nassicurazione attrezzature;\ndeposito sicuro"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  17. ALLEGATI
# ═══════════════════════════════════════════════════════════
doc.add_heading('17. Allegati', level=1)

doc.add_paragraph(
    "Si elencano di seguito gli allegati da predisporre e da inserire nella documentazione "
    "per la partecipazione ai bandi:"
)

add_numbered(doc, "Curricula vitae degli istruttori federali FICK/FICT")
add_numbered(doc, "Copia dei brevetti/qualifiche FICK/FICT degli istruttori")
add_numbered(doc, "Visura camerale della Sib.Bellini Turist Service di Leli Emanuela")
add_numbered(doc, "Documentazione della concessione dell'area lacuale")
add_numbered(doc, "Statuto dell'ASD (da redigere)")
add_numbered(doc, "Atto costitutivo dell'ASD (da redigere)")
add_numbered(doc, "Accordo di partenariato tra ASD e Sib.Bellini Turist Service di Leli Emanuela")
add_numbered(doc, "Preventivi attrezzature da almeno 3 fornitori")
add_numbered(doc, "Preventivo assicurativo")
add_numbered(doc, "Planimetria dell'area con indicazione punto imbarco/sbarco e deposito")
add_numbered(doc, "Documentazione fotografica del sito")
add_numbered(doc, "Programma didattico dei corsi (conforme FICK/FICT)")
add_numbered(doc, "Protocolli di sicurezza")
add_numbered(doc, "Piano di marketing dettagliato")
add_numbered(doc, "Dichiarazione di rispetto del principio DNSH")
add_numbered(doc, "Lettera di intenti delle strutture ricettive partner")
add_numbered(doc, "Eventuali lettere di supporto istituzionale (Comune, Pro Loco, Ente Parco)")

doc.add_paragraph()
doc.add_paragraph()

# ── Note finali ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fine del documento —")
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Documento predisposto come base per la partecipazione a bandi pubblici.\n"
    "Da integrare con dati specifici, allegati e personalizzazioni richieste\n"
    "dal singolo bando di riferimento."
)
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ═══════════════════════════════════════════════════════════
#  SALVA
# ═══════════════════════════════════════════════════════════
doc.save(OUTPUT_FILE)
print(f"DOCX generato: {OUTPUT_FILE}")
